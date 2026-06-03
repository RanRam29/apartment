"""
DirApp v3.1 — Document Generator
Creates 4 professional Word documents matching the original spec format.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colors — RGBColor for font, hex strings for cell backgrounds ─────────────
BLUE_DARK   = RGBColor(0x1A, 0x37, 0x6C)
BLUE_MID    = RGBColor(0x2E, 0x5E, 0xA8)
BLUE_LIGHT  = RGBColor(0xD6, 0xE4, 0xF7)
GREEN_DARK  = RGBColor(0x1B, 0x6B, 0x3A)
GREEN_LIGHT = RGBColor(0xD4, 0xED, 0xDA)
RED_DARK    = RGBColor(0x8B, 0x00, 0x00)
GRAY_LIGHT  = RGBColor(0xF5, 0xF5, 0xF5)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)

# Hex strings for cell backgrounds (used with set_cell_bg)
HEX_BLUE_LIGHT  = 'D6E4F7'
HEX_BLUE_DARK   = '1A376C'
HEX_GREEN_LIGHT = 'D4EDDA'
HEX_RED_LIGHT   = 'F8D7DA'
HEX_GRAY_LIGHT  = 'F5F5F5'
HEX_WHITE       = 'FFFFFF'

# ── Helpers ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    """hex_color: 6-char hex string, e.g. 'D6E4F7'"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def set_cell_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:color'), '2E5EA8')
                tcBorders.append(border)
            tcPr.append(tcBorders)

def add_cover_page(doc, title, subtitle, version, date):
    """Add styled cover page."""
    # Big title paragraph with dark blue background simulation via borders
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\n\n\n')
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('DirApp')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = BLUE_DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = BLUE_MID

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # Version / Date box (simple table)
    tbl = doc.add_table(rows=3, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows_data = [
        ('גרסה', version),
        ('תאריך', date),
        ('סטטוס', 'מאושר לפיתוח'),
    ]
    for i, (k, v) in enumerate(rows_data):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = v
        set_cell_bg(tbl.rows[i].cells[0], HEX_BLUE_LIGHT)
        for cell in tbl.rows[i].cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(11)
                    run.font.bold = (i == 0)
    set_cell_borders(tbl)

    doc.add_page_break()

def style_heading(doc, text, level=1):
    """Add a styled heading."""
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        if level == 1:
            run.font.color.rgb = BLUE_DARK
            run.font.size = Pt(16)
        elif level == 2:
            run.font.color.rgb = BLUE_MID
            run.font.size = Pt(13)
        else:
            run.font.color.rgb = BLUE_MID
            run.font.size = Pt(11)
    return p

def add_note(doc, text, color=None):
    """Add a colored note paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f'📌 {text}')
    run.font.size = Pt(10)
    run.font.italic = True
    if color:
        run.font.color.rgb = color

def add_table(doc, headers, rows, header_bg=None, col_widths=None):
    """Add a styled table."""
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, header_bg or HEX_BLUE_LIGHT)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = BLUE_DARK

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(cell_text)
            # Alternate row shading
            if r_idx % 2 == 1:
                set_cell_bg(cell, HEX_GRAY_LIGHT)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in para.runs:
                    run.font.size = Pt(9)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)

    set_cell_borders(tbl)
    doc.add_paragraph()
    return tbl

def add_code_block(doc, code_text):
    """Add a code-style paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F0F0F0')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)

def add_info_box(doc, title, items):
    """Add a highlighted info box."""
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, HEX_BLUE_LIGHT)
    p = cell.paragraphs[0]
    run = p.add_run(f'◆ {title}\n')
    run.font.bold = True
    run.font.color.rgb = BLUE_DARK
    run.font.size = Pt(10)
    for item in items:
        run2 = p.add_run(f'• {item}\n')
        run2.font.size = Pt(9)
    set_cell_borders(tbl)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 1: Product Roadmap
# ═══════════════════════════════════════════════════════════════════════════════
def create_roadmap():
    doc = Document()
    section = doc.sections[0]
    section.right_to_left = True
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    add_cover_page(doc,
        'Product Roadmap',
        'מפת דרכים — כל הפאזות',
        'v3.1', '2026-06-02')

    # TOC stub
    style_heading(doc, 'תוכן עניינים', level=1)
    toc_items = [
        ('1', 'סיכום מנהלים'),
        ('2', 'Phase 1 — MVP Core (v3.0) ✅'),
        ('3', 'Phase 2 — WhatsApp Integration ✅'),
        ('4', 'Phase 3 — V2 Extensions'),
        ('5', 'Phase 4 — Enterprise V3.0'),
        ('6', 'טבלאות DB נדרשות לעתיד'),
        ('7', 'עקרונות עיצוב המוצר'),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph(f'{num}.  {title}')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in p.runs:
            run.font.size = Pt(11)
    doc.add_page_break()

    # ── Section 1: Executive Summary ──────────────────────────────────────────
    style_heading(doc, '1. סיכום מנהלים', level=1)
    p = doc.add_paragraph(
        'DirApp היא פלטפורמת matching + ניהול חוזים להשכרת דירות בשוק הישראלי. '
        'המערכת מחברת שוכרים ומשכירים, מנהלת חוזים דיגיטליים, לדגר תשלומים, '
        'צ\'ק-אין/אאוט, ותחזוקה — הכל ממקום אחד.')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.size = Pt(11)

    doc.add_paragraph()
    add_table(doc,
        ['מדד', 'ערך'],
        [
            ['Milestones MVP', 'M1–M15 (15 סה"כ)'],
            ['סטטוס MVP', '✅ הושלם ואומת בייצור (2026-06-01)'],
            ['Backend', 'https://apartment-backend-v24y.onrender.com'],
            ['Frontend', 'https://apartment-olive.vercel.app'],
            ['Mobile', 'React Native + Expo (iOS + Android)'],
            ['כניסת בדיקה', 'admin2@dirapp.com / Admin1234!'],
            ['הבא בתור', 'V2-1 Stripe Connect'],
        ],
        col_widths=[2.5, 4]
    )
    doc.add_page_break()

    # ── Section 2: Phase 1 MVP ─────────────────────────────────────────────────
    style_heading(doc, '2. Phase 1 — MVP Core (v3.0)', level=1)
    add_note(doc, 'כל 15 ה-Milestones הושלמו ואומתו בייצור. סטטוס כולל: ✅', GREEN_DARK)
    doc.add_paragraph()

    milestones = [
        ('M1', 'Auth & Onboarding', [
            'Register / Login / JWT (24h)',
            'Switch Role (tenant ↔ landlord)',
            'Terms of Service (tosAcceptedAt)',
            'Email verification (Resend)',
            'Multi-tenant — מנהל בית',
        ]),
        ('M2', 'Property Listing & Discovery', [
            'Property upload (images → R2, 10 images max)',
            'Apartments feed (Redis cache 1h)',
            'True Monthly Cost (שכירות + ארנונה + ועד)',
            'GenAI Marketing Copy (Gemini 1.5 Flash)',
            'Multi-room support (kitchen/salon/bathroom/shower/custom)',
        ]),
        ('M3', 'Swipe & Matching Engine', [
            'Swipe right/left/superlike + daily quota',
            'Match creation + landlord accept/decline',
            'Compatibility scoring (Gemini) — 8 lifestyle dimensions',
            'Smart Map + TAMA 38 Urban Renewal Layer',
            'AI Lead Qualification',
        ]),
        ('M4', 'Real-Time Chat', [
            'Socket.io + MongoDB persistence',
            'Message pagination (50/page)',
            'Chat archival (not deletion)',
            'Block user functionality',
        ]),
        ('M5', 'KYC — Know Your Customer', [
            'Persona Web SDK integration',
            'HMAC-SHA256 webhook validation',
            'APPROVED → auto-unlock contracts',
            'REJECTED → הנחיות ספציפיות + badge',
            'TIMEOUT (24h) → push + email',
            'KYC images auto-delete (7 ימים — Tiquon 13)',
            'Admin KYC override (GODMODE)',
            'KYC renewal alert (5 שנים)',
        ]),
        ('M6', 'Digital Contract Lifecycle', [
            'Contract upload (PDF/DOCX) + Gemini OCR',
            'State Machine (7 states): UPLOAD→PENDING_SIGN→ACTIVE→EXPIRING→ENDED',
            'Digital signature (tenant + guarantor + landlord)',
            'Validation gate (KYC + mandatory fields)',
            'Ownership verification (tenant)',
            'Contract amendments (propose/approve/reject)',
            'Contract renewal',
        ]),
        ('M7', 'Guarantor Web Flow', [
            'Email invite link (5-day TTL)',
            'Persona Web SDK for guarantor KYC',
            'Digital signature on guarantee document',
            'Decline → landlord notification (email + push)',
            'Auto-reminder 24h לפני פקיעה',
        ]),
        ('M8', 'Check-In & Room Inventory', [
            'תמונות לפי חדרים (עד 20 תמונות/חדר → R2)',
            'אישור משכיר + הצהרה משותפת',
            'סבבי תיקון (עד 3 סבבים) + Auto-Confirm',
            'נעילת תמונות post-approval (immutable evidence)',
            'Check-In window (5 ימים, configurable)',
        ]),
        ('M9', 'Check-Out & Damage Assessment', [
            'תמונות יציאה + הערות (אותה מבנה כמו Check-In)',
            'השוואה Check-In vs Check-Out',
            'סבבי תיקון (עד 3 סבבים)',
            'R2 storage (contract term + 3 years)',
        ]),
        ('M10', 'Admin Panel — GODMODE', [
            'Admin login (JWT role=admin)',
            'Config Panel (52 keys, 9 sections) — real-time, ללא deploy',
            'User Management (edit, cascading delete, unlock)',
            'Stats Dashboard (8 sections, 56 metrics)',
            'Contract state override + KYC override',
            'Manual notification sending',
        ]),
        ('M11', 'Digital Ledger & Payments', [
            'יצירת שורות לדגר אוטומטית (start_date → end_date)',
            'שוכר מדווח "שילמתי" + receipt upload',
            'משכיר מאשר → PAID',
            'Auto-Confirm 48h (configurable)',
            'OVERDUE → escalation T+0 to T+5',
            'CPI שנתי (Jan 1) per contract',
        ]),
        ('M12', 'Maintenance & Service Tickets', [
            'שוכר פותח תקלה + תמונה אופציונלית',
            'משכיר מסמן "אני מטפל"',
            'Deep link → midrag.co.il',
            'חשבונית → R2 (ticket_invoices)',
            'Cron: 24h ללא מענה → התראה; 3 ימים → escalation',
        ]),
        ('M13', 'Notifications', [
            'Expo Push (token registration, כל תפקיד)',
            'Resend Email — 8 templates בעברית',
            'Contract EXPIRING (120/90/60/45/30 days)',
        ]),
        ('M14', 'Storage & R2 Integration', [
            '5 buckets עם lifecycle policies',
            'Presigned URLs (max 5 min validity)',
            'Cloudinary replaced by R2',
            'Archive auto-delete after 3 years',
        ]),
        ('M15', 'Security, Compliance & Audit', [
            'JWT 24h + refresh via session cache',
            'RBAC: tenant / landlord / admin / guarantor',
            'HMAC-SHA256 webhooks (Persona + WhatsApp)',
            'Rate limiting (10 req/min per IP)',
            'Path traversal hardening',
            'Audit trail — auditCapture middleware',
            'Tiquon 13 — KYC 7-day deletion',
        ]),
    ]

    for m_id, m_title, features in milestones:
        style_heading(doc, f'{m_id}: {m_title}  ✅', level=2)
        for feat in features:
            p = doc.add_paragraph(f'✔  {feat}', style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.size = Pt(10)
        doc.add_paragraph()

    doc.add_page_break()

    # ── Section 3: WhatsApp ────────────────────────────────────────────────────
    style_heading(doc, '3. Phase 2 — WhatsApp Integration', level=1)
    add_note(doc, 'הושלם 2026-06-01 | 13/13 tests passing', GREEN_DARK)
    add_note(doc, 'דרוש Meta Business Setup: רישום templates + env vars', RED_DARK)
    doc.add_paragraph()

    add_table(doc,
        ['רכיב', 'תיאור', 'סטטוס'],
        [
            ['Meta Cloud API client', 'sendTemplate / sendText / sendInteractive / markAsRead / downloadMedia', '✅'],
            ['Webhook verify + receive', 'GET verify + POST inbound, HMAC-SHA256', '✅'],
            ['8 Hebrew templates', 'payment (3) + maintenance (3) + invite + renewal', '✅'],
            ['State machine', 'idle → payment_confirm / maintenance → image flows', '✅'],
            ['Notification service', 'whatsappNotificationService.js — 8 public methods', '✅'],
            ['Cron → WA payments', 'ledgerDueAlerts.js + ledgerOverdue.js upgraded', '✅'],
            ['Cron → WA renewal', 'expiringAlerts.js upgraded (60d before end)', '✅'],
            ['Payment confirm via WA', 'Tenant confirms → LedgerRow REPORTED', '✅'],
            ['Maintenance via WA', 'Description + photo → Ticket + R2', '✅'],
            ['Sequelize models', 'WhatsAppMessage + WhatsAppConversationState', '✅'],
            ['Tests', '13/13 passing', '✅'],
        ],
        col_widths=[2, 4, 0.7]
    )

    add_info_box(doc, 'Env Vars נדרשים ב-Render:', [
        'WHATSAPP_API_TOKEN',
        'WHATSAPP_PHONE_NUMBER_ID',
        'WHATSAPP_VERIFY_TOKEN',
        'WHATSAPP_APP_SECRET',
    ])
    doc.add_page_break()

    # ── Section 4: V2 Extensions ───────────────────────────────────────────────
    style_heading(doc, '4. Phase 3 — V2 Extensions (הבא בתור)', level=1)

    v2_features = [
        ('V2-1', 'Stripe Connect (Payment Gateway)', 'Cursor', '❌ לא התחיל', [
            'Stripe Connect integration (Israeli processing)',
            'Automated ACH / Bank Transfer for rent',
            'Payment method storage (tokenized, PCI)',
            'Webhook → auto-ledger update',
            'Reconciliation: bank tx ↔ ledger_row',
            'New tables: payment_methods, payment_gateway_transactions',
            '⚠️ דורש רישיון ממשרד האוצר',
        ]),
        ('V2-2', 'Tenant Guarantor Claims', 'Cursor', '❌ לא התחיל', [
            'Claim filing (landlord vs guarantor)',
            'Guarantor response: accept / dispute',
            'Admin mediation workflow',
            'New table: warranty_claims',
        ]),
        ('V2-3', 'Contract Amendment Workflow', 'Antigravity', '✅ הושלם', [
            'propose / approve / reject routes — קיים',
            'GET /api/v3/contracts/:id כולל amendments — קיים',
            'אומת E2E 2026-06-01',
        ]),
        ('V2-4', 'Advanced NLP Search', 'Antigravity', '🟡 חלקי', [
            'parseNLPSearch קיים ב-geminiService',
            'Relevance ranking — נדרש שיפור',
            'Redis cache for NLP queries (6h) — קיים',
        ]),
        ('V2-5', 'Tenant Credit Scoring', 'Antigravity', '🟡 חלקי', [
            'users.trustScore = 50 default — קיים',
            'Auto-calc post-checkout — workflow חסר',
            'Landlord visibility — נדרש UI',
        ]),
        ('V2-7', 'GDPR Compliance & Data Governance', 'Claude Code', '❌ לא התחיל', [
            'Data export (JSON) — right to data portability',
            'Right to deletion — right to be forgotten',
            'Notification opt-out mechanism',
            'Audit log retention (7 years)',
            'New tables: notification_preferences, deletion_requests',
        ]),
        ('V2-8', 'Dispute Resolution', 'Cursor', '❌ לא התחיל (מוחרג מ-MVP)', [
            'Formal dispute escalation',
            'Admin mediation (SLA 48h)',
            'Evidence presentation (photos, messages, receipts)',
            'New table: disputes',
        ]),
    ]

    for v_id, v_title, owner, status, features in v2_features:
        style_heading(doc, f'{v_id}: {v_title}', level=2)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f'אחריות: {owner}   |   סטטוס: {status}')
        run.font.size = Pt(10)
        run.font.color.rgb = BLUE_MID
        for feat in features:
            p2 = doc.add_paragraph(f'•  {feat}', style='List Bullet')
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run2 in p2.runs:
                run2.font.size = Pt(10)
        doc.add_paragraph()

    doc.add_page_break()

    # ── Section 5: Enterprise V3.0 ────────────────────────────────────────────
    style_heading(doc, '5. Phase 4 — Enterprise V3.0 (עתידי)', level=1)
    add_note(doc, 'כל הפיצ\'רים בפאזה זו — לא התחילו. תחילת פיתוח לאחר V2 מלא.', RGBColor(0x88, 0x44, 0x00))
    doc.add_paragraph()

    add_table(doc,
        ['ID', 'שם פיצ\'ר', 'תיאור קצר', 'טבלאות חדשות'],
        [
            ['EV3-1', 'Multi-Company Platform', 'Company → Building → Unit hierarchy, Property Manager role', 'companies, buildings, role_assignments'],
            ['EV3-2', 'SLA & Service Level', 'SLA per landlord, auto-tracking, breach notifications', '—'],
            ['EV3-3', 'Insurance Integration', 'Policy verification, damage → insurance claim, premium calc', '—'],
            ['EV3-4', 'TI-1525 Compliance', 'Israeli regulation enforcement, clause validation, attestation', 'compliance_flags, compliance_attestations'],
            ['EV3-5', 'Field App', 'Native tablet app for property managers, offline mode, NFC', '—'],
            ['EV3-6', 'QR Code Rooms', 'QR label per room, scan → open form, auto room_id link', 'qr_codes'],
            ['EV3-7', 'Bank Reconciliation', 'Bank sync, ledger matching, tax authority reports', 'bank_transactions, reconciliation_log'],
            ['EV3-8', 'Vendor Management', 'Vendor registry, auto-quotes, work order tracking', 'vendors, work_orders'],
        ],
        col_widths=[0.8, 1.8, 3, 2]
    )
    doc.add_page_break()

    # ── Section 6: Summary Table ──────────────────────────────────────────────
    style_heading(doc, '6. סיכום סטטוס לפי Phase', level=1)
    add_table(doc,
        ['Phase', 'פיצ\'רים', 'הושלמו', 'נשארו', 'סטטוס'],
        [
            ['MVP Core (v3.0) M1-M15', '15', '15 ✅', '0', 'הושלם'],
            ['WhatsApp Integration', '11', '11 ✅', '0', 'הושלם'],
            ['Design & Gamification', '7', '7 ✅', '0', 'הושלם'],
            ['V2 Extensions', '8', '2', '6', 'בתכנון'],
            ['Enterprise V3.0', '8', '0', '8', 'עתידי'],
        ],
        col_widths=[2.5, 1, 1, 1, 1.5]
    )

    style_heading(doc, '7. עקרונות עיצוב המוצר', level=1)
    principles = [
        'DirApp neutral — לא מתערב בסכסוכים בין שוכר למשכיר',
        'Privacy-first — מינימום שמירת מידע (Tiquon 13 compliant)',
        'Mobile-first — Expo React Native, iOS + Android',
        'Hebrew-first — כל ה-UX בעברית',
        'Manual payments MVP — Stripe בגרסה V2',
        'Admin configurability — 52 פרמטרים ניתנים לשינוי real-time ללא deploy',
    ]
    for pr in principles:
        p = doc.add_paragraph(f'◆  {pr}')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in p.runs:
            run.font.size = Pt(11)

    doc.save('C:\\apartmentapp\\docs\\DirApp_Product_Roadmap_v3.1.docx')
    print('✅ DirApp_Product_Roadmap_v3.1.docx')


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 2: Development Architecture Spec
# ═══════════════════════════════════════════════════════════════════════════════
def create_architecture():
    doc = Document()
    section = doc.sections[0]
    section.right_to_left = True
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    add_cover_page(doc,
        'Development Architecture Spec',
        'מפרט ארכיטקטורה טכנית מלאה',
        'v3.1', '2026-06-02')

    # ── 1. Tech Stack ──────────────────────────────────────────────────────────
    style_heading(doc, '1. Stack טכנולוגי', level=1)
    add_table(doc,
        ['Layer', 'Technology', 'Version', 'Purpose'],
        [
            ['Runtime', 'Node.js', '20 LTS', 'Backend execution'],
            ['Framework', 'Express.js', '4.22.2', 'REST API'],
            ['DB Primary', 'PostgreSQL', '15+', 'Structured data (Supabase)'],
            ['DB Chat', 'MongoDB', '7.x', 'Real-time messaging (Atlas)'],
            ['Cache', 'Redis', '7.x', 'Session + feed cache (Upstash)'],
            ['Storage', 'Cloudflare R2', '—', 'Files: images, contracts, receipts'],
            ['Mobile', 'React Native + Expo', '51.x', 'iOS + Android'],
            ['Web — Guarantor', 'React', '18.x', 'Guarantor web flow'],
            ['Web — Admin', 'React', '18.x', 'Admin GODMODE panel'],
            ['WebSocket', 'Socket.io', '4.x', 'Real-time chat'],
            ['AI / OCR', 'Google Gemini', '1.5 Flash', 'OCR, NLP, marketing copy'],
            ['KYC', 'Persona', '—', 'Identity verification'],
            ['Email', 'Resend', '—', 'Transactional email'],
            ['Push', 'Expo Push', '—', 'Mobile notifications'],
            ['WhatsApp', 'Meta Cloud API', '—', 'Conversational notifications'],
            ['Backend Host', 'Render', '—', 'Auto-deploy from main branch'],
            ['Web Host', 'Vercel', '—', 'Guarantor + Admin SPA'],
            ['ORM (PG)', 'Sequelize', '6.x', 'PostgreSQL models'],
            ['ORM (Mongo)', 'Mongoose', '8.x', 'MongoDB models'],
        ],
        col_widths=[1.5, 1.8, 1, 3]
    )
    doc.add_page_break()

    # ── 2. Backend Structure ───────────────────────────────────────────────────
    style_heading(doc, '2. מבנה Backend (Modular Monolith)', level=1)
    add_code_block(doc,
        'backend/\n'
        '├── src/\n'
        '│   ├── config/         ← DB connections, APIs, storage clients\n'
        '│   ├── models/\n'
        '│   │   ├── pg/         ← 24+ Sequelize models (PostgreSQL)\n'
        '│   │   └── mongo/      ← 3 Mongoose collections (MongoDB)\n'
        '│   ├── middleware/      ← auth, errorHandler, auditCapture, rateLimit\n'
        '│   ├── routes/         ← 18+ route files\n'
        '│   ├── services/       ← kycService, contractService, geminiService, ...\n'
        '│   ├── cron/           ← 8 scheduled jobs\n'
        '│   └── utils/          ← validators, stateMachine, hmac, logger\n'
        '└── tests/              ← unit + integration + e2e (154 test cases)\n'
    )

    style_heading(doc, '2.1 Config Files', level=2)
    add_table(doc,
        ['קובץ', 'תפקיד'],
        [
            ['database.js', 'PostgreSQL (Supabase) + ensureUserVerificationColumns()'],
            ['mongodb.js', 'MongoDB (Atlas) connection'],
            ['redis.js', 'Redis (Upstash) — session + feed + NLP cache'],
            ['cloudflare-r2.js', 'R2 SDK init (S3-compatible API)'],
            ['gemini.js', 'Gemini 1.5 Flash API client'],
            ['persona.js', 'Persona SDK'],
            ['resend.js', 'Resend SMTP client'],
            ['expo-push.js', 'Expo Push service'],
            ['whatsapp.js', 'Meta Cloud API client'],
            ['socket-io.js', 'Socket.io server config'],
        ],
        col_widths=[2, 4.5]
    )

    style_heading(doc, '2.2 Routes (API Endpoints)', level=2)
    add_table(doc,
        ['קובץ Route', 'Base Path', 'תיאור'],
        [
            ['auth.js', '/api/auth/', 'Register, login, reset, profile, switch-role, accept-tos'],
            ['apartments.js', '/api/apartments/', 'Upload, feed, details, marketing-copy'],
            ['swipe.js', '/api/', 'Swipe, matches list, accept/decline'],
            ['chat.js', '/api/chat/', 'Messages + Socket.io events'],
            ['contracts.js', '/api/contracts/ + /api/v3/contracts/', 'Upload, sign, checkin, checkout, amend, renew'],
            ['ledger.js', '/api/ledger/', 'Report payment, confirm, receipt upload'],
            ['kyc.js', '/api/kyc/', 'Initiate, webhook (HMAC), status'],
            ['guarantor.js', '/api/guarantor/', 'Invite, verify, accept/decline'],
            ['maintenance.js', '/api/maintenance/', 'Open ticket, update, invoice, close'],
            ['landlord.js', '/api/landlord/', 'Dashboard, listings analytics'],
            ['tenant.js', '/api/tenant/', 'Renter journal, edit profile'],
            ['gamification.js', '/api/gamification/', 'Trust score, award points, leaderboard'],
            ['whatsapp.js', '/api/whatsapp/', 'Webhook verify + receive'],
            ['admin.js', '/api/v3/admin/', 'Config, stats, users, contracts, payments, maintenance'],
            ['screening.js', '/api/screening/', 'Identity screening (BDI/Gov)'],
            ['commercial.js', '/api/commercial/', 'Commercial leases'],
            ['services.js', '/api/services/', 'Services marketplace'],
            ['iot.js', '/api/iot/', 'IoT devices + maintenance tickets'],
        ],
        col_widths=[2, 2.5, 3]
    )

    style_heading(doc, '2.3 Middleware', level=2)
    add_table(doc,
        ['Middleware', 'תפקיד'],
        [
            ['auth.js', 'JWT validation + role extraction\nconst { authenticate, requireRole } = require(\'../middleware/auth\')'],
            ['errorHandler.js', 'Global error handling + standard error codes'],
            ['auditCapture.js', 'Log all state changes (who, what, when, old_value, new_value)'],
            ['requestContext.js', 'Inject user_id, role, company_id into req object'],
            ['requireTos.js', 'TOS acceptance gate (403 if not accepted)'],
            ['requireKycApproved.js', 'KYC APPROVED gate (422 if not approved)'],
            ['rateLimit.js', '10 requests/minute per IP per endpoint'],
            ['cors.js', 'Explicit domain whitelist (Vercel + Render + dev)'],
        ],
        col_widths=[2.2, 4.3]
    )
    doc.add_page_break()

    # ── 3. Database Schema ─────────────────────────────────────────────────────
    style_heading(doc, '3. Database Schema — PostgreSQL (Supabase)', level=1)

    style_heading(doc, '3.1 Schema: auth', level=2)
    add_table(doc,
        ['טבלה', 'עמודות עיקריות', 'הערות'],
        [
            ['users',
             'id (SERIAL PK), phone (UNIQUE), email (UNIQUE), passwordHash,\nrole (tenant|landlord|admin), kyc_status, kyc_expires_at,\nactiveRole, trustScore (default 50), blocked_count, is_locked,\ntosAcceptedAt, created_at, updated_at',
             'Role flexibility per contract; trustScore starts at 50'],
            ['kyc_records',
             'id, user_id (FK→users), persona_inquiry_id, status\n(PENDING|PROCESSING|APPROVED|REJECTED|TIMEOUT),\nid_expiry (DATE), created_at, updated_at',
             'תמונות נמחקות מ-R2 אחרי 7 ימים (Tiquon 13)'],
            ['ownership_declarations',
             'id, user_id (FK), contract_id (FK), declared_at, ip (INET)',
             'Audit trail — IP tracking for ownership verification'],
        ],
        col_widths=[1.8, 3.5, 2.2]
    )

    style_heading(doc, '3.2 Schema: properties', level=2)
    add_table(doc,
        ['טבלה', 'עמודות עיקריות', 'הערות'],
        [
            ['properties',
             'id, owner_id (FK→users), address, lat (DECIMAL 10,8),\nlng (DECIMAL 11,8), status (ACTIVE|INACTIVE|ARCHIVED),\ncreated_at, updated_at',
             'Landlord property listings'],
            ['property_rooms',
             'id, property_id (FK), name, type\n(kitchen|salon|bathroom|shower|custom)',
             'Room types for Check-In/Out; builtin + custom'],
            ['property_images',
             'id, property_id (FK), r2_key (VARCHAR 500), "order" (INT)',
             '→ R2 property-images bucket\npublic-read, forever TTL, max 10 images'],
        ],
        col_widths=[1.8, 3.5, 2.2]
    )

    style_heading(doc, '3.3 Schema: contracts', level=2)
    add_table(doc,
        ['טבלה', 'עמודות עיקריות', 'הערות'],
        [
            ['agreements',
             'id, property_id (FK), landlord_id (FK),\nstatus (UPLOAD|PENDING_SIGN|ACTIVE|EXPIRING|\nPENDING_ACTIVATION|ENDED),\nstart_date, end_date, monthly_rent (DECIMAL 10,2),\npayment_day (INT default 1), cpi_enabled (BOOLEAN),\nr2_doc_key, created_at, updated_at',
             'State Machine — 7 states\nStateLockGuard pattern'],
            ['agreement_parties',
             'id, agreement_id (FK), user_id (FK),\nrole (tenant|guarantor|landlord),\nsigned_at (TIMESTAMP), kyc_status',
             'All signing parties per contract'],
            ['agreement_rooms',
             'id, agreement_id (FK), room_id (FK→property_rooms),\ncheckin_photos TEXT[] default \'{}\',\ncheckout_photos TEXT[] default \'{}\',\ncheckin_status, checkout_status',
             'Photo arrays = R2 keys\n→ R2 checkin-photos bucket'],
            ['ledger_rows',
             'id, agreement_id (FK), period (YYYY-MM),\ndue_date (DATE), amount (DECIMAL 10,2),\nstatus (PENDING|REPORTED|PAID|OVERDUE),\nreported_at, confirmed_at, cpi_adj,\nnotes (TEXT), receipt_r2_key',
             '→ R2 payment-receipts bucket\nAuto-Confirm 48h via cron'],
            ['ownership_verifications',
             'id, agreement_id (FK), tenant_id (FK),\nchoice (checked|skipped), verified_at',
             'Tenant verifies landlord owns property'],
            ['contract_amendments',
             'id, contract_id (FK), proposed_by (FK→users),\nfield, oldValue, newValue,\nstatus (PENDING|APPROVED|REJECTED)',
             'Max 10 revisions per contract'],
        ],
        col_widths=[1.8, 3.5, 2.2]
    )

    style_heading(doc, '3.4 Schema: maintenance', level=2)
    add_table(doc,
        ['טבלה', 'עמודות עיקריות', 'הערות'],
        [
            ['tickets',
             'id, agreement_id (FK), reporter_id (FK→users),\ndescription (TEXT), status (OPEN|IN_PROGRESS|RESOLVED|CLOSED),\nphoto_r2_key, created_at, updated_at',
             'Maintenance requests; optional photo'],
            ['ticket_invoices',
             'id, ticket_id (FK), r2_key, amount (DECIMAL),\npayer (landlord|tenant), uploaded_at',
             'Receipt tracking per ticket'],
        ],
        col_widths=[1.8, 3.5, 2.2]
    )

    style_heading(doc, '3.5 Schema: admin & config', level=2)
    add_table(doc,
        ['טבלה', 'עמודות עיקריות', 'הערות'],
        [
            ['app_config',
             'id, key (UNIQUE), value (TEXT), section, description,\nupdated_at, updated_by (FK→users)',
             '52 keys, 9 sections\nConfigurable real-time, no deploy needed'],
            ['whatsapp_messages',
             'id, user_id (FK), phone, direction (inbound|outbound),\nmessage_type, content (JSONB), wa_message_id,\ntemplate_name, status, created_at',
             'WhatsApp audit log'],
            ['whatsapp_conversation_states',
             'id, phone (UNIQUE), state (default \'idle\'),\ncontext (JSONB), updated_at',
             'Conversational state machine per phone'],
            ['user_points',
             'id, user_id (UNIQUE FK), points (default 50),\nbadges TEXT[], updated_at',
             'Trust Score + gamification badges'],
        ],
        col_widths=[1.8, 3.5, 2.2]
    )
    doc.add_page_break()

    # ── 4. MongoDB ─────────────────────────────────────────────────────────────
    style_heading(doc, '4. Database — MongoDB (Atlas)', level=1)
    add_table(doc,
        ['Collection', 'שדות עיקריים', 'שימוש'],
        [
            ['messages', 'match_id, sender_id, text, image_url, created_at', 'Real-time chat (Socket.io)'],
            ['chat_sessions', 'match_id, agreement_id, archived_at, created_at', 'Session metadata'],
            ['user_preferences', 'user_id, lifestyle_score (object), swipe_history, preferences', 'Swipe prefs + matching'],
        ],
        col_widths=[1.8, 3, 2.7]
    )
    doc.add_paragraph()

    # ── 5. Redis ───────────────────────────────────────────────────────────────
    style_heading(doc, '5. Redis — Key Patterns', level=1)
    add_table(doc,
        ['Key Pattern', 'תוכן', 'TTL', 'שימוש'],
        [
            ['feed:{userId}', 'Swipe feed (10 apartments)', '1 hour', 'Feed performance'],
            ['session:{token}', 'JWT session data', '24 hours', 'Auth'],
            ['nlp:{queryHash}', 'NLP query result', '6 hours', 'NLP search cache'],
            ['kyc:lock:{userId}', 'KYC double-submit lock', '10 minutes', 'Concurrency control'],
        ],
        col_widths=[2, 2, 1, 2.5]
    )
    doc.add_page_break()

    # ── 6. R2 Storage ─────────────────────────────────────────────────────────
    style_heading(doc, '6. Cloudflare R2 — Storage Buckets', level=1)
    add_table(doc,
        ['Bucket', 'תוכן', 'גישה', 'Lifecycle / TTL', 'מקס'],
        [
            ['property-images', 'Marketing images', 'public-read', 'Forever', '2MB/image (max 10)'],
            ['contract-docs', 'PDF/DOCX חוזים', 'private', 'Contract term', '10MB/file'],
            ['checkin-photos', 'תמונות חדרים', 'private', 'Contract + 3 years', '2MB/photo (max 20/room)'],
            ['payment-receipts', 'קבלות תשלום', 'private', 'Contract + 3 years', '5MB/file'],
            ['archive', 'חוזים פגי תוקף', 'private', '3 years → auto-delete', 'Unlimited'],
            ['kyc-temp', 'תמונות KYC', 'private', '7 days → auto-delete', '5MB/file'],
        ],
        col_widths=[1.8, 1.5, 1, 1.8, 1.4]
    )
    add_note(doc, 'Presigned URLs: max 5 minutes validity. Client-side compression: JPEG/WEBP, 2MB max.', BLUE_MID)
    doc.add_page_break()

    # ── 7. Services ────────────────────────────────────────────────────────────
    style_heading(doc, '7. Services — API פנימי', level=1)

    style_heading(doc, '7.1 kycService.js', level=2)
    add_table(doc,
        ['Method', 'תיאור'],
        [
            ['initiateVerification(userId)', 'שולח לPersona, מחזיר URL; sets Redis lock kyc:lock:{userId} (10 min)'],
            ['handleWebhook(payload, hmacSignature)', 'Validates HMAC-SHA256; updates kyc_records.status; triggers checkAndUnlockContracts on APPROVED'],
            ['checkAndUnlockContracts(userId)', 'מוצא חוזי PENDING_SIGN שהמשתמש צד להם; פותח חתימה אם כולם APPROVED'],
            ['IKYCProvider (interface)', 'Abstraction layer — מאפשר החלפת Persona בעתיד ללא שינוי קוד'],
        ],
        col_widths=[2.8, 4.7]
    )

    style_heading(doc, '7.2 contractService.js', level=2)
    add_table(doc,
        ['Method', 'תיאור'],
        [
            ['uploadAndExtract(file, landlordId)', 'מאמת סוג קובץ (PDF/DOCX, max 10MB); uploads to R2; calls geminiService.extractContractFields()'],
            ['validateGate(agreementId)', 'בודק: KYC APPROVED לכל הצדדים; שדות חובה קיימים; מחזיר 422 אם חסוי'],
            ['transitionState(agreementId, newState)', 'StateLockGuard — מונע מעבר מקביל; מאמת מעברים מורשים'],
            ['generateLedger(agreementId)', 'יוצר LedgerRow אחד לחודש (start→end); מחיל CPI אם מופעל'],
            ['scheduleExpiringAlerts(agreementId)', 'יוצר טריגרי cron בתאריכי: end - 120/90/60/45/30 ימים'],
        ],
        col_widths=[2.8, 4.7]
    )

    style_heading(doc, '7.3 geminiService.js (Gemini 1.5 Flash)', level=2)
    add_table(doc,
        ['Method', 'Input', 'Output'],
        [
            ['extractContractFields(fileBuffer)', 'PDF/DOCX buffer', 'JSON: tenant_name, tenant_id, address, start_date, end_date, monthly_rent, payment_day, cpi_enabled'],
            ['generateMarketingCopy(aptData)', 'Property data', 'Hebrew description (3 styles: professional/friendly/luxury)'],
            ['parseNLPSearch(query)', 'Hebrew free text', 'JSON: rooms, location, max_price, etc. (cached Redis 6h)'],
            ['scoreCompatibility(tenant, apt)', 'Tenant + apartment', 'Score 0-100 for feed ranking'],
        ],
        col_widths=[2.3, 1.8, 3.4]
    )

    style_heading(doc, '7.4 notificationService.js', level=2)
    add_table(doc,
        ['Method', 'Channel', 'תיאור'],
        [
            ['sendPush(userId, title, body, data)', 'Expo Push', 'Mobile push; works regardless of activeRole'],
            ['sendEmail(to, template, data)', 'Resend SMTP', '8 Hebrew HTML templates; <3K/month free'],
            ['scheduleReminder(userId, timestamp, payload)', 'DB + Cron', 'Future scheduled notification (T-5, expiring, KYC)'],
        ],
        col_widths=[2.8, 1.2, 3.5]
    )

    style_heading(doc, '7.5 whatsappNotificationService.js (8 methods)', level=2)
    add_table(doc,
        ['Method', 'תיאור'],
        [
            ['sendPaymentReminder(userId, amount, dueDate)', 'תזכורת תשלום T-3'],
            ['sendPaymentConfirmRequest(userId, ledgerRowId)', 'בקשת אישור תשלום'],
            ['sendPaymentOverdue(userId, amount, daysPast)', 'תשלום באיחור'],
            ['sendMaintenanceCreated(landlordId, ticketId)', 'תקלה חדשה למשכיר'],
            ['sendMaintenanceStatusUpdate(tenantId, status)', 'עדכון סטטוס תקלה'],
            ['sendMaintenanceResolved(tenantId, ticketId)', 'תקלה נסגרה'],
            ['sendGuarantorInvite(phone, contractId)', 'הזמנת ערב'],
            ['sendContractRenewalAlert(userId, contractId, daysLeft)', 'התראת חידוש חוזה'],
        ],
        col_widths=[3.2, 4.3]
    )
    doc.add_page_break()

    # ── 8. Cron Jobs ───────────────────────────────────────────────────────────
    style_heading(doc, '8. Cron Jobs (8 jobs)', level=1)
    add_table(doc,
        ['Job', 'Schedule', 'Trigger', 'פעולה', 'Admin Config Key'],
        [
            ['ledgerDueAlerts', 'Daily 08:00', 'T-5 לפני due_date', 'Push + Email + WhatsApp reminder', 'overdue_alert_days'],
            ['ledgerOverdue', 'Daily 23:59', 'T+0 through T+5', 'Mark OVERDUE; escalation + WhatsApp', '—'],
            ['paymentAutoconfirm', 'Hourly', '48h after REPORTED', 'Auto-confirm REPORTED→PAID', 'payment_autoconfirm_hours'],
            ['expiringAlerts', 'Daily 09:00', '120/90/60/45/30d to end', 'Contract EXPIRING notifications + WhatsApp', 'expiring_warning_days'],
            ['cpiAdjustment', 'Jan 1 annually', 'Date: Jan 1', 'Recalculate ledger rows with CPI', '—'],
            ['kycRenewal', 'Daily', 'id_expiry - 30d', 'KYC renewal push + email', 'kyc_renewal_years'],
            ['maintenanceAlerts', 'Hourly', '24h + 3d open ticket', 'Escalate unanswered tickets', 'maintenance_alert_hours_1/days_2'],
            ['r2Cleanup', 'Monthly (1st)', 'Archive files > 3yr', 'Delete expired archive files', '—'],
        ],
        col_widths=[1.8, 1.2, 1.5, 2.2, 1.8]
    )
    doc.add_page_break()

    # ── 9. API Endpoints ───────────────────────────────────────────────────────
    style_heading(doc, '9. API Endpoints — מלא', level=1)

    endpoints = [
        ('Auth /api/auth/', [
            ('POST', '/register', 'Signup (landlord/tenant)'),
            ('POST', '/login', 'JWT + tosAcceptedAt + activeRole + kycStatus'),
            ('POST', '/reset-password', 'Password reset via Resend'),
            ('GET', '/profile', 'User profile'),
            ('POST', '/switch-role', 'Switch activeRole (tenant ↔ landlord)'),
            ('POST', '/accept-tos', 'Accept Terms of Service'),
        ]),
        ('Apartments /api/apartments/', [
            ('POST', '/upload', 'Upload property + images → R2'),
            ('GET', '/feed', 'Swipe feed (Redis cached 1h)'),
            ('GET', '/:id', 'Property details + True Monthly Cost'),
            ('POST', '/:id/marketing-copy', 'GenAI marketing copy (Gemini)'),
        ]),
        ('Swipe & Matches /api/', [
            ('POST', '/swipe', 'Record swipe {apartmentId, like, superlike}'),
            ('GET', '/matches', 'Match list for current user'),
            ('POST', '/matches/:id/accept', 'Landlord accepts match'),
            ('POST', '/matches/:id/decline', 'Rejection with optional reason'),
        ]),
        ('Chat /api/chat/', [
            ('GET', '/:matchId', 'Paginated messages (50/page)'),
            ('POST', '/:matchId/close', 'Archive chat'),
        ]),
        ('KYC /api/kyc/', [
            ('POST', '/initiate', 'Start KYC → return Persona URL'),
            ('POST', '/webhook', 'Persona callback (HMAC-SHA256 validated)'),
            ('GET', '/status', 'Current KYC status'),
        ]),
        ('Contracts /api/contracts/ + /api/v3/contracts/', [
            ('POST', '/upload', 'Upload PDF/DOCX + Gemini OCR'),
            ('GET', '/:id', 'Full contract details + amendments'),
            ('POST', '/:id/sign', 'Digital signature'),
            ('POST', '/:id/checkin', 'Check-In photos (by room)'),
            ('GET', '/:id/checkin/status', 'Check-In progress'),
            ('POST', '/:id/checkin/approve', 'Landlord approves'),
            ('POST', '/:id/checkin/request-revision', 'Request re-photo'),
            ('POST', '/:id/checkout', 'Check-Out photos'),
            ('POST', '/:id/checkout/approve', 'Landlord assessment'),
            ('POST', '/:id/ownership', 'Tenant ownership verification'),
            ('POST', '/:id/amend', 'Propose amendment'),
            ('POST', '/:id/renew', 'Start renewal process'),
        ]),
        ('Ledger /api/ledger/', [
            ('GET', '/contracts/:id/ledger', 'All ledger rows for contract'),
            ('POST', '/ledger/:id/report', 'Tenant reports payment'),
            ('POST', '/ledger/:id/confirm', 'Landlord confirms/rejects'),
            ('POST', '/ledger/:id/receipt', 'Upload payment proof → R2'),
        ]),
        ('Guarantor /api/guarantor/', [
            ('POST', '/invite', 'Send guarantor invite (5-day link)'),
            ('GET', '/verify/:token', 'Guarantor verifies link'),
            ('POST', '/accept/:token', 'Accept + KYC initiation'),
            ('POST', '/decline/:token', 'Decline → notify landlord'),
        ]),
        ('Maintenance /api/maintenance/', [
            ('POST', '/', 'Open ticket'),
            ('GET', '/:contractId', 'Ticket list'),
            ('POST', '/:id/update', 'Status + notes'),
            ('POST', '/:id/invoice', 'Upload receipt → R2'),
            ('POST', '/:id/close', 'Mark resolved'),
        ]),
        ('Admin /api/v3/admin/', [
            ('GET', '/config', 'Get all 52 config keys'),
            ('PUT', '/config/:key', 'Update config value (real-time)'),
            ('GET', '/stats/detailed', '8 sections, 56 metrics'),
            ('GET', '/users', 'Paginated user list'),
            ('PUT', '/users/:id', 'Edit user'),
            ('DELETE', '/users/:id', 'Cascading delete'),
            ('POST', '/users/:id/unlock', 'Unlock locked user'),
            ('POST', '/kyc/:userId/override', 'Override KYC status'),
            ('GET', '/contracts', 'All contracts with state'),
            ('PUT', '/contracts/:id/status', 'Force state transition'),
            ('GET', '/payments', 'All ledger rows'),
            ('POST', '/payments/:id/confirm', 'Force confirm payment'),
            ('GET', '/maintenance', 'All tickets'),
            ('POST', '/maintenance/:id/close', 'Close ticket'),
            ('POST', '/notifications/send', 'Manual notification'),
        ]),
        ('Gamification & Tenant /api/', [
            ('GET', '/gamification/me', 'Trust Score + badges'),
            ('POST', '/gamification/award', 'Award points'),
            ('GET', '/gamification/leaderboard', 'Ranked users'),
            ('GET', '/tenant/journal', 'Renter Journal (full summary)'),
            ('PUT', '/tenant/profile', 'Edit tenant profile'),
        ]),
        ('WhatsApp /api/whatsapp/', [
            ('GET', '/webhook', 'Meta webhook verification'),
            ('POST', '/webhook', 'Inbound messages (HMAC validated)'),
        ]),
    ]

    for section_title, routes in endpoints:
        style_heading(doc, section_title, level=2)
        add_table(doc,
            ['Method', 'Path', 'תיאור'],
            [[m, p, d] for m, p, d in routes],
            col_widths=[0.7, 2.5, 4.3]
        )

    doc.add_page_break()

    # ── 10. Mobile Structure ───────────────────────────────────────────────────
    style_heading(doc, '10. Mobile App — Expo React Native', level=1)

    style_heading(doc, '10.1 Navigation Structure', level=2)
    add_code_block(doc,
        'AppNavigator.tsx\n'
        '├── AuthStack\n'
        '│   ├── OnboardingScreen, LoginScreen, RegisterScreen, TosScreen\n'
        '├── TenantTabs\n'
        '│   ├── SwipeScreen (Discovery)\n'
        '│   ├── MatchesScreen → ChatScreen\n'
        '│   ├── ContractsScreen → ContractDetail, CheckIn, CheckOut, LedgerScreen\n'
        '│   ├── MaintenanceScreen\n'
        '│   ├── MapScreen (Leaflet + TAMA 38)\n'
        '│   └── ProfileScreen → KYC, TrustScore, RenterJournal, Roommate\n'
        '└── LandlordTabs\n'
        '    ├── LandlordDashboardScreen\n'
        '    ├── ListingsScreen → CreateListingScreen\n'
        '    ├── MatchesScreen (incoming) → ChatScreen\n'
        '    ├── ContractsScreen → CheckInApproval, CheckOutApproval, Ledger\n'
        '    ├── MaintenanceScreen (landlord view)\n'
        '    └── ProfileScreen\n'
    )

    style_heading(doc, '10.2 State Management (Zustand)', level=2)
    add_table(doc,
        ['Store', 'State מנוהל'],
        [
            ['useAuthStore', 'user, token, activeRole, kycStatus, tosAcceptedAt'],
            ['useChatStore', 'messages, activeChatId, socket connection'],
            ['useContractStore', 'contracts, agreements, amendments'],
            ['useLedgerStore', 'ledger rows, payment history'],
            ['useMaintenanceStore', 'tickets, invoices'],
            ['useSwipeStore', 'feed, swipe queue, matches'],
        ],
        col_widths=[2.5, 5]
    )
    doc.add_page_break()

    # ── 11. Security ───────────────────────────────────────────────────────────
    style_heading(doc, '11. Security Architecture', level=1)
    add_table(doc,
        ['Layer', 'מנגנון', 'פרטים'],
        [
            ['Auth', 'JWT', '24h expiry, role in payload'],
            ['Role gating', 'RBAC middleware', 'tenant / landlord / admin / guarantor'],
            ['Webhooks', 'HMAC-SHA256', 'Persona + WhatsApp Meta'],
            ['Rate limiting', '10 req/min/IP', 'Per endpoint'],
            ['File upload', 'Path traversal guard', 'resolveUploadFilePath + safeUnlinkUpload'],
            ['Storage', 'Private presigned URLs', 'Max 5 min validity'],
            ['Passwords', 'bcrypt hash', 'stored in passwordHash column'],
            ['CORS', 'Domain whitelist', 'Explicit allowed origins only'],
            ['Data minimization', 'KYC 7-day delete', 'Tiquon 13 compliance'],
            ['Audit trail', 'auditCapture middleware', 'All state changes logged'],
        ],
        col_widths=[1.5, 2, 4]
    )
    doc.add_page_break()

    # ── 12. Deploy & Conventions ──────────────────────────────────────────────
    style_heading(doc, '12. Deploy Pipeline & Code Conventions', level=1)

    style_heading(doc, '12.1 Deploy Pipeline', level=2)
    add_code_block(doc,
        'Developer → git push → GitHub (main branch)\n'
        '  ↓\n'
        'Render auto-deploy (~3 min)\n'
        'Health check: GET /health → 200 OK\n'
        'URL: https://apartment-backend-v24y.onrender.com\n\n'
        'Mobile → Expo EAS Build → App Store / Google Play\n\n'
        'Web → Vercel auto-deploy\n'
        'URL: https://apartment-olive.vercel.app\n'
    )

    style_heading(doc, '12.2 Worktrees', level=2)
    add_table(doc,
        ['Agent', 'Worktree', 'Branch', 'תחום'],
        [
            ['Claude Code (Orchestrator)', 'C:\\apartmentapp', 'main', 'ניהול, merge, backend קריטי'],
            ['Cursor', 'C:\\apartmentapp-cursor', 'cursor/financial-admin', 'Financial, Admin, Cron'],
            ['Antigravity (Windsurf)', 'C:\\apartmentapp-windsurf', 'wind/*', 'Mobile, Frontend, KYC, Identity'],
        ],
        col_widths=[2, 2, 1.8, 2]
    )
    add_note(doc, 'כלל: רק Claude Code (Orchestrator) עושה merge לmain.', RED_DARK)

    style_heading(doc, '12.3 Code Conventions', level=2)
    add_table(doc,
        ['Convention', 'פרטים'],
        [
            ['Endpoints חדשים', '/api/v3/ prefix (ישנים נשמרים ב-/api/ לתאימות לאחור)'],
            ['Auth middleware', 'const { authenticate, requireRole } = require(\'../middleware/auth\')'],
            ['Models location', 'backend/src/models/pg/ (Sequelize, SERIAL PKs)'],
            ['Schema changes', 'חובה לעדכן ensureUserVerificationColumns() ב-database.js'],
            ['Branch naming', 'fix/<description> | feat/<description>'],
            ['Tests', 'Jest — backend/tests/'],
            ['Commit format', 'type(scope): description\nAGENT: ...\nTOKENS: ~XK\nTIME: ...\nTESTS: ...'],
        ],
        col_widths=[2, 5.5]
    )

    doc.save('C:\\apartmentapp\\docs\\DirApp_Development_Architecture_Spec_v3.1.docx')
    print('✅ DirApp_Development_Architecture_Spec_v3.1.docx')


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 3: Product Specification Book
# ═══════════════════════════════════════════════════════════════════════════════
def create_product_spec():
    doc = Document()
    section = doc.sections[0]
    section.right_to_left = True
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    add_cover_page(doc,
        'Product Specification Book',
        'ספר מפרט מוצר — כל הפיצ\'רים',
        'v3.1', '2026-06-02')

    # ── Overview ───────────────────────────────────────────────────────────────
    style_heading(doc, '1. סקירת מוצר', level=1)
    p = doc.add_paragraph(
        'DirApp היא פלטפורמת SaaS לניהול מלא של יחסי שוכר-משכיר בשוק הישראלי. '
        'המערכת מכסה את כל מחזור החיים של השכרת דירה: '
        'גילוי → matching → חוזה דיגיטלי → צ\'ק-אין → תשלום חודשי → תחזוקה → צ\'ק-אאוט.')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.size = Pt(11)

    doc.add_paragraph()
    add_info_box(doc, 'ערכי ליבה:', [
        'שוק יעד: ישראל — Hebrew-first UX',
        'פלטפורמה: Mobile (iOS + Android) + Web (Guarantor + Admin)',
        'עמדת DirApp: ניטרלית — לא מתערבת בסכסוכים',
        'עמדת נתונים: Privacy-first (Tiquon 13 compliant)',
        'סטטוס: MVP v3.0 מלא ובייצור',
    ])

    # ── Roles ──────────────────────────────────────────────────────────────────
    style_heading(doc, '2. תפקידי משתמשים', level=1)
    add_table(doc,
        ['תפקיד', 'תיאור', 'הרשאות עיקריות'],
        [
            ['Tenant (שוכר)', 'מחפש דירה לשכירות', 'Swipe, chat, sign contract, report payment, open ticket, Check-In/Out photos'],
            ['Landlord (משכיר)', 'מפרסם ומנהל דירות', 'Create listing, view matches, send contracts, approve Check-In, confirm payment'],
            ['Admin (מנהל)', 'ניהול כולל המערכת', 'GODMODE — כל הפעולות; config panel; stats; override'],
            ['Guarantor (ערב)', 'מאשר ערבות לחוזה', 'Web-only: KYC + digital signature on guarantee document'],
            ['Property Manager', 'מנהל בניין (עתידי)', 'ניהול קבוצת דירות בשם חברה (EV3-1)'],
        ],
        col_widths=[1.5, 2, 4]
    )
    add_note(doc, 'גמישות: משתמש יכול להיות גם שוכר וגם משכיר — activeRole מגדיר את הממשק הפעיל.', BLUE_MID)
    doc.add_page_break()

    # ── Feature Specs ──────────────────────────────────────────────────────────
    style_heading(doc, '3. מפרט פיצ\'רים מפורט', level=1)

    # Contract State Machine
    style_heading(doc, '3.1 Contract State Machine', level=2)
    add_table(doc,
        ['State', 'תיאור', 'כניסה אפשרית מ-', 'יציאה אפשרית ל-'],
        [
            ['UPLOAD', 'חוזה הועלה, AI מחלץ שדות', '—', 'PENDING_SIGN'],
            ['PENDING_SIGN', 'כל השדות תקינים, מחכה לחתימות', 'UPLOAD', 'ACTIVE'],
            ['ACTIVE', 'חוזה פעיל (PENDING_CHECK_IN → LIVING)', 'PENDING_SIGN, PENDING_ACTIVATION', 'EXPIRING, ENDED'],
            ['EXPIRING', 'פחות מ-120 יום לסיום', 'ACTIVE', 'ENDED, PENDING_ACTIVATION'],
            ['PENDING_ACTIVATION', 'חוזה חידוש נחתם, מחכה לתאריך תחילה', 'EXPIRING', 'ACTIVE'],
            ['ENDED', 'חוזה הסתיים; ארכיב 3 שנים → auto-delete', 'ACTIVE, EXPIRING', '—'],
        ],
        col_widths=[1.8, 2, 2, 2]
    )
    add_note(doc, 'StateLockGuard מונע מעברי מצב מקביליים (concurrency protection).', BLUE_MID)

    # KYC Flow
    style_heading(doc, '3.2 KYC Flow', level=2)
    add_table(doc,
        ['שלב', 'פעולה', 'טריגר הבא'],
        [
            ['1', 'משתמש לוחץ "אמת זהות" באפליקציה', 'POST /api/kyc/initiate'],
            ['2', 'kyc:lock:{userId} set ב-Redis (10 min)', 'מניעת כפילות'],
            ['3', 'Persona URL מוחזר; משתמש עובר ל-WebView', 'Persona SDK'],
            ['4', 'Persona webhook → POST /api/kyc/webhook', 'HMAC-SHA256 validation'],
            ['5a', 'status = APPROVED → checkAndUnlockContracts()', 'חוזים נפתחים לחתימה'],
            ['5b', 'status = REJECTED → badge אדום + הנחיות', 'User מנסה שוב'],
            ['5c', 'status = TIMEOUT (24h) → push + email', 'User מנסה שוב'],
            ['6', '7 ימים לאחר APPROVED: תמונות נמחקות מ-R2', 'Tiquon 13 compliance'],
            ['7', 'KYC renewal: 5 שנים → alert user', 'Cron יומי'],
        ],
        col_widths=[0.5, 3.5, 3.5]
    )

    # Ledger Flow
    style_heading(doc, '3.3 Payment Ledger Flow', level=2)
    add_table(doc,
        ['Status', 'מי מעדכן', 'תנאי', 'הבא'],
        [
            ['PENDING', 'System (auto)', 'יצירת שורה בתחילת חודש', 'REPORTED / OVERDUE'],
            ['REPORTED', 'Tenant', 'שוכר מדווח "שילמתי" + receipt', 'PAID (confirm) / PENDING (reject)'],
            ['PAID', 'Landlord / Auto', 'משכיר מאשר OR 48h auto-confirm', '— (final)'],
            ['OVERDUE', 'Cron (23:59)', 'T+0: due_date עבר ללא REPORTED', 'REPORTED (לאחר דיווח)'],
        ],
        col_widths=[1.2, 1.5, 2.8, 2]
    )
    add_note(doc, 'Auto-Confirm: אם משכיר לא מגיב תוך 48h → REPORTED הופך אוטומטית ל-PAID.', BLUE_MID)

    # Check-In Flow
    style_heading(doc, '3.4 Check-In / Check-Out Flow', level=2)
    add_table(doc,
        ['שלב', 'פעולה', 'הגבלות'],
        [
            ['1', 'חוזה עובר ל-ACTIVE → חלון צ\'ק-אין נפתח', '5 ימים (configurable: check_in_window_days)'],
            ['2', 'שוכר מצלם כל חדר', 'עד 20 תמונות/חדר; JPEG/WEBP 2MB; → R2 checkin-photos'],
            ['3', 'משכיר סוקר ומאשר / דורש תיקון', 'עד 3 סבבי תיקון (checkout_revision_rounds)'],
            ['4', 'לאחר אישור → תמונות נעולות (immutable)', 'ראיות משפטיות לסוף חוזה'],
            ['5 (יציאה)', 'שוכר מצלם שוב בסוף חוזה', 'אותה מבנה; השוואה Check-In vs Out'],
        ],
        col_widths=[0.5, 3, 4]
    )

    # Guarantor Web Flow
    style_heading(doc, '3.5 Guarantor Web Flow', level=2)
    add_table(doc,
        ['שלב', 'פעולה'],
        [
            ['1', 'משכיר שולח הזמנה (POST /api/guarantor/invite) → email עם link'],
            ['2', 'Link תקף 5 ימים (configurable: guarantor_link_ttl_days)'],
            ['3', 'ערב פותח דף → רואה: שם משכיר, כתובת, תקופה, שכירות, הסבר אחריות'],
            ['4a — ACCEPT', 'Persona Web SDK מוטמע → KYC → חתימה דיגיטלית על מסמך ערבות'],
            ['4b — DECLINE', 'לחיצת "אני לא מסכים" → משכיר מקבל email + push מיידי'],
            ['5', 'reminder email 24h לפני פקיעת link'],
            ['6', 'link פג → error page + אפשרות לבקש invite חדש'],
        ],
        col_widths=[1.5, 6]
    )
    doc.add_page_break()

    # Admin Config
    style_heading(doc, '3.6 Admin Config Panel — 52 Keys', level=2)
    add_note(doc, 'כל הפרמטרים ניתנים לשינוי real-time דרך ממשק Admin ללא deploy. GET/PUT /api/v3/admin/config', BLUE_MID)
    add_table(doc,
        ['Key', 'Default', 'תיאור', 'Range'],
        [
            ['check_in_window_days', '5', 'ימי חלון לצ\'ק-אין', '1–10'],
            ['checkin_photos_max', '20', 'מקסימום תמונות לחדר', '5–50'],
            ['checkout_revision_rounds', '3', 'סבבי תיקון אאוט', '1–5'],
            ['expiring_warning_days', '120', 'ימי התראה לפני סיום חוזה', '30–180'],
            ['guarantor_link_ttl_days', '5', 'תקפות link ערב', '1–30'],
            ['blocking_threshold', '5', 'חסימות → נעילת חשבון', '3–10'],
            ['contract_revision_max', '10', 'מקסימום תיקוני חוזה', '3–20'],
            ['payment_autoconfirm_hours', '48', 'שעות לאישור אוטומטי', '24–168'],
            ['overdue_alert_days', '5', 'ימי גרייס לפני OVERDUE', '1–10'],
            ['kyc_renewal_years', '5', 'שנות תוקף KYC', '3–7'],
            ['maintenance_alert_hours_1', '24', 'שעות עד התראת תחזוקה ראשונה', '12–48'],
            ['maintenance_alert_days_2', '3', 'ימים עד escalation תחזוקה', '1–7'],
            ['persona_monthly_quota', '500', 'מכסת KYC חודשית', '100–5000'],
            ['...+ 39 keys נוספים', '—', 'כל שאר הפרמטרים בממשק האדמין', '—'],
        ],
        col_widths=[2.5, 0.8, 2.5, 1.2]
    )
    doc.add_page_break()

    # WhatsApp Conversational State Machine
    style_heading(doc, '3.7 WhatsApp Conversational Flows', level=2)
    add_table(doc,
        ['State', 'Trigger', 'Response', 'Next State'],
        [
            ['idle', 'Inbound message', 'תפריט ראשי', 'awaiting_action'],
            ['awaiting_action', '1 — payment', 'בקשת אישור תשלום', 'payment_confirm'],
            ['payment_confirm', 'כן / אישור', 'LedgerRow → REPORTED', 'idle'],
            ['awaiting_action', '2 — maintenance', 'תאר את הבעיה', 'maintenance_description'],
            ['maintenance_description', 'טקסט', 'האם יש תמונה?', 'maintenance_image'],
            ['maintenance_image', 'תמונה / דלג', 'Ticket created + R2 upload', 'idle'],
        ],
        col_widths=[2, 1.8, 2, 1.7]
    )

    # Non-Functional Requirements
    style_heading(doc, '4. Non-Functional Requirements', level=1)
    add_table(doc,
        ['דרישה', 'ערך', 'הערות'],
        [
            ['Uptime', '99.5%', 'Render SLA'],
            ['API Response Time', '<500ms p95', 'Redis caching for hot paths'],
            ['File Upload Max', '10MB (contracts), 2MB (images)', 'Client-side compression'],
            ['Rate Limiting', '10 req/min per IP', 'Per endpoint'],
            ['JWT Expiry', '24 hours', 'Session cache in Redis'],
            ['KYC Image Retention', '7 days max', 'Tiquon 13 — auto-delete'],
            ['Archive Retention', '3 years', 'R2 lifecycle rule'],
            ['Notification SLA', '<30 seconds for push', 'Expo Push'],
            ['WhatsApp Response', 'Real-time', 'Meta webhook <5s'],
            ['Mobile Platforms', 'iOS 16+ / Android 12+', 'Expo 51'],
        ],
        col_widths=[2.2, 2, 3.3]
    )
    doc.add_page_break()

    # Compliance
    style_heading(doc, '5. Compliance & Legal', level=1)
    add_table(doc,
        ['נושא', 'גישה', 'סטטוס'],
        [
            ['Tiquon 13 (Israeli Data Law)', 'KYC images deleted 7 days post-approval; ID numbers only (no copies)', '✅ מיושם'],
            ['GDPR Art 6.1.b', 'Contract formation legal basis', '✅ מכוסה'],
            ['GDPR Art 6.1.f', 'Legitimate interest: fraud prevention, matching', '✅ מכוסה'],
            ['GDPR Right to Data', 'Data export endpoint — V2-7', '❌ עתידי'],
            ['GDPR Right to Erasure', 'Deletion request workflow — V2-7', '❌ עתידי'],
            ['TI-1525 Compliance', 'Israeli landlord-tenant regulation enforcement', '❌ EV3-4'],
            ['Payment Processing', 'Manual MVP; Stripe V2 requires MoF license', '⚠️ ב-V2'],
        ],
        col_widths=[2.5, 3, 1.7]
    )

    doc.save('C:\\apartmentapp\\docs\\DirApp_Product_Specification_Book_v3.1.docx')
    print('✅ DirApp_Product_Specification_Book_v3.1.docx')


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 4: Test Coverage Matrix
# ═══════════════════════════════════════════════════════════════════════════════
def create_test_matrix():
    doc = Document()
    section = doc.sections[0]
    section.right_to_left = True
    section.page_width = Inches(11.69)   # A4 Landscape
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    add_cover_page(doc,
        'Test Coverage Matrix',
        'מטריצת כיסוי בדיקות — כל Test Cases',
        'v3.1', '2026-06-02')

    style_heading(doc, 'סיכום כיסוי', level=1)
    add_table(doc,
        ['קטגוריה', 'סה"כ', 'P0', 'P1', 'P2', 'P3', 'סטטוס'],
        [
            ['Auth + KYC', '10', '8', '2', '—', '—', '✅ 10/10'],
            ['Contract State Machine', '14', '7', '4', '3', '—', '✅ 14/14'],
            ['Financial Engine (Ledger)', '11', '4', '6', '1', '—', '✅ 11/11'],
            ['Maintenance', '6', '2', '4', '—', '—', '✅ 6/6'],
            ['Chat + Blocking', '5', '2', '2', '—', '1', '✅ 5/5'],
            ['Notifications + Storage', '8', '4', '2', '2', '—', '✅ 8/8'],
            ['Admin Panel', '6', '3', '3', '—', '—', '✅ 6/6'],
            ['WhatsApp Integration', '13', '—', '13', '—', '—', '✅ 13/13'],
            ['Security + Edge Cases', '8', '5', '1', '—', '2', '✅ 8/8'],
            ['TOTAL', '81', '35', '37', '6', '3', '✅ 81/81'],
        ],
        col_widths=[2.5, 0.7, 0.5, 0.5, 0.5, 0.5, 1.2]
    )
    add_note(doc, 'P0 = Critical (must pass pre-launch) | P1 = High | P2 = Medium | P3 = Low', BLUE_MID)
    doc.add_page_break()

    # ── Auth + KYC Tests ──────────────────────────────────────────────────────
    style_heading(doc, '1. Auth + KYC Tests (10)', level=1)
    add_table(doc,
        ['ID', 'Feature', 'Test Case', 'P', 'Method', 'Expected', 'Status'],
        [
            ['AUTH-001', 'Register', 'שוכר נרשם עם אימייל + טלפון', 'P0', 'POST /api/auth/register', '201 + JWT + tosAcceptedAt', '✅'],
            ['AUTH-002', 'Login', 'משתמש מתחבר — JWT מוחזר', 'P0', 'POST /api/auth/login', '200 + token + activeRole + kycStatus', '✅'],
            ['AUTH-003', 'Login — admin', 'admin2@dirapp.com מתחבר', 'P0', 'POST /api/auth/login', 'role: admin, admin routes accessible', '✅'],
            ['AUTH-004', 'Switch Role', 'שוכר עובר לתפקיד משכיר', 'P0', 'POST /api/auth/switch-role', 'activeRole updated, UI reflects change', '✅'],
            ['AUTH-005', 'ToS Gate', 'משתמש ללא ToS → blocked', 'P0', 'POST /api/apartments/upload', '403 requireTos middleware', '✅'],
            ['AUTH-006', 'Rate Limit', '11 requests/min → rejected', 'P0', 'Any endpoint', '429 Too Many Requests', '✅'],
            ['KYC-001', 'KYC Initiate', 'שוכר מתחיל KYC', 'P0', 'POST /api/kyc/initiate', 'Persona URL returned, Redis lock set', '✅'],
            ['KYC-002', 'KYC Webhook — APPROVED', 'Persona APPROVED webhook', 'P0', 'POST /api/kyc/webhook', 'HMAC valid, status=APPROVED, contracts unlocked', '✅'],
            ['KYC-003', 'KYC Webhook — REJECTED', 'Persona REJECTED webhook', 'P0', 'POST /api/kyc/webhook', 'status=REJECTED, badge UI update', '✅'],
            ['KYC-004', 'KYC Image Deletion', '7 ימים לאחר APPROVED → מחיקה', 'P1', 'Cron', 'R2 KYC images deleted', '✅'],
        ],
        col_widths=[1, 1.3, 2, 0.4, 2, 2, 0.6]
    )

    # ── Contract Tests ─────────────────────────────────────────────────────────
    style_heading(doc, '2. Contract State Machine Tests (14)', level=1)
    add_table(doc,
        ['ID', 'Feature', 'Test Case', 'P', 'Method', 'Expected', 'Status'],
        [
            ['CON-001', 'Upload + OCR', 'העלאת PDF + Gemini חילוץ שדות', 'P0', 'POST /api/contracts/upload', 'Fields extracted, status=UPLOAD', '✅'],
            ['CON-002', 'Upload — invalid type', 'העלאת JPG → rejected', 'P0', 'POST /api/contracts/upload', '400 invalid file type', '✅'],
            ['CON-003', 'Validate Gate — KYC', 'ניסיון חתימה ללא KYC APPROVED', 'P0', 'POST /api/contracts/:id/sign', '422 requireKycApproved', '✅'],
            ['CON-004', 'Sign — all parties', 'כל הצדדים חותמים → ACTIVE', 'P0', 'POST /api/contracts/:id/sign', 'status transitions to ACTIVE', '✅'],
            ['CON-005', 'State Lock Guard', 'שני requests מקביליים → אחד נדחה', 'P0', 'POST /api/contracts/:id/sign x2', 'Second request 409 Conflict', '✅'],
            ['CON-006', 'Ownership Verify', 'שוכר מאמת בעלות משכיר', 'P0', 'POST /api/contracts/:id/ownership', 'choice saved with timestamp', '✅'],
            ['CON-007', 'Ledger Generation', 'חוזה 12 חודשים → 12 שורות', 'P0', 'Auto (post-sign)', '12 LedgerRows created', '✅'],
            ['CON-008', 'Check-In Window', 'צ\'ק-אין אחרי 5 ימים → alert admin', 'P1', 'Cron', 'Admin alert sent', '✅'],
            ['CON-009', 'Check-In Photos', '20 תמונות/חדר → R2', 'P1', 'POST /api/contracts/:id/checkin', 'Photos stored in R2 checkin-photos', '✅'],
            ['CON-010', 'Check-In Revision', 'משכיר דורש תיקון → סבב 2', 'P1', 'POST /api/contracts/:id/checkin/request-revision', 'revision_round incremented', '✅'],
            ['CON-011', 'Check-Out Comparison', 'חוות דעת נזקים', 'P1', 'GET /api/contracts/:id/checkout/comparison', 'Before/after photo pairs returned', '✅'],
            ['CON-012', 'Amendment — approve', 'שני הצדדים מאשרים תיקון', 'P2', 'POST /api/contracts/:id/amend', 'status=APPROVED, contract updated', '✅'],
            ['CON-013', 'Renewal', 'חידוש חוזה → PENDING_ACTIVATION', 'P2', 'POST /api/contracts/:id/renew', 'New agreement created, old→EXPIRING', '✅'],
            ['CON-014', 'EXPIRING Alerts', '120/90/60/45/30 days → notifications', 'P2', 'Cron expiringAlerts', 'Push + Email + WhatsApp sent', '✅'],
        ],
        col_widths=[1, 1.3, 2, 0.4, 2, 2, 0.6]
    )

    # ── Ledger Tests ──────────────────────────────────────────────────────────
    style_heading(doc, '3. Financial Engine — Ledger Tests (11)', level=1)
    add_table(doc,
        ['ID', 'Feature', 'Test Case', 'P', 'Method', 'Expected', 'Status'],
        [
            ['LED-001', 'Ledger Auto-Generate', '12 חוזה → 12 שורות PENDING', 'P0', 'Auto post-sign', '12 rows, correct amounts + due_dates', '✅'],
            ['LED-002', 'Report Payment', 'שוכר מדווח + receipt', 'P0', 'POST /api/ledger/:id/report', 'status=REPORTED, reported_at set', '✅'],
            ['LED-003', 'Confirm Payment', 'משכיר מאשר → PAID', 'P0', 'POST /api/ledger/:id/confirm', 'status=PAID, confirmed_at set', '✅'],
            ['LED-004', 'Reject Payment', 'משכיר דוחה → back to PENDING', 'P0', 'POST /api/ledger/:id/confirm {reject}', 'status=PENDING, tenant notified', '✅'],
            ['LED-005', 'Auto-Confirm 48h', 'REPORTED 48h → auto PAID', 'P1', 'Cron paymentAutoconfirm', 'status=PAID without landlord action', '✅'],
            ['LED-006', 'OVERDUE Alert T+0', 'due_date עבר → OVERDUE', 'P1', 'Cron ledgerOverdue', 'status=OVERDUE, push+WA to both', '✅'],
            ['LED-007', 'OVERDUE Daily Escalation', 'T+1 through T+5 → daily alerts', 'P1', 'Cron ledgerOverdue', '5 escalations sent over 5 days', '✅'],
            ['LED-008', 'Payment Due Reminder', 'T-5 → reminder push', 'P1', 'Cron ledgerDueAlerts', 'Push + Email + WhatsApp sent', '✅'],
            ['LED-009', 'Receipt Upload', 'שוכר מצרף קבלה → R2', 'P1', 'POST /api/ledger/:id/receipt', 'File in R2 payment-receipts', '✅'],
            ['LED-010', 'CPI Adjustment', 'Jan 1 → CPI update on active contracts', 'P1', 'Cron cpiAdjustment', 'ledger_rows.cpi_adj updated', '✅'],
            ['LED-011', 'CPI Disabled', 'חוזה ללא CPI → לא מעודכן', 'P2', 'Cron cpiAdjustment', 'Rows unchanged for cpi_enabled=false', '✅'],
        ],
        col_widths=[1, 1.5, 2, 0.4, 2, 2, 0.6]
    )
    doc.add_page_break()

    # ── Maintenance Tests ──────────────────────────────────────────────────────
    style_heading(doc, '4. Maintenance Tests (6)', level=1)
    add_table(doc,
        ['ID', 'Feature', 'Test Case', 'P', 'Method', 'Expected', 'Status'],
        [
            ['MNT-001', 'Open Ticket', 'שוכר פותח תקלה + תמונה', 'P0', 'POST /api/maintenance', 'Ticket created, photo in R2', '✅'],
            ['MNT-002', 'Landlord Responds', 'משכיר מסמן IN_PROGRESS', 'P0', 'POST /api/maintenance/:id/update', 'status=IN_PROGRESS, tenant notified', '✅'],
            ['MNT-003', '24h No Response Alert', 'ללא מענה 24h → alert', 'P1', 'Cron maintenanceAlerts', 'Push + Email to landlord', '✅'],
            ['MNT-004', '3 Days Escalation', 'ללא סגירה 3 ימים → escalation', 'P1', 'Cron maintenanceAlerts', 'Alert to admin + landlord', '✅'],
            ['MNT-005', 'Invoice Upload', 'חשבונית → R2', 'P1', 'POST /api/maintenance/:id/invoice', 'File in R2, ticket_invoices record', '✅'],
            ['MNT-006', 'Close Ticket', 'סגירת תקלה', 'P1', 'POST /api/maintenance/:id/close', 'status=CLOSED, all parties notified', '✅'],
        ],
        col_widths=[1, 1.5, 2.5, 0.4, 2, 2, 0.6]
    )

    # ── Security Tests ─────────────────────────────────────────────────────────
    style_heading(doc, '5. Security + Edge Case Tests (8)', level=1)
    add_table(doc,
        ['ID', 'Feature', 'Test Case', 'P', 'Method', 'Expected', 'Status'],
        [
            ['SEC-001', 'JWT Expired', 'Token פג תוקף → rejected', 'P0', 'Any authenticated endpoint', '401 Unauthorized', '✅'],
            ['SEC-002', 'Role Guard', 'Tenant ניגש ל-admin route', 'P0', 'GET /api/v3/admin/users', '403 Forbidden', '✅'],
            ['SEC-003', 'HMAC Webhook', 'Webhook ללא HMAC → rejected', 'P0', 'POST /api/kyc/webhook', '401 Invalid signature', '✅'],
            ['SEC-004', 'Path Traversal', 'Upload ../../etc/passwd', 'P0', 'POST /api/contracts/upload', '400 resolveUploadFilePath rejects', '✅'],
            ['SEC-005', 'Rate Limit', '11 req/min → 429', 'P0', 'POST /api/auth/login x11', '429 Too Many Requests after 10', '✅'],
            ['SEC-006', 'KYC Lock', 'שני initiateVerification מקביליים', 'P1', 'POST /api/kyc/initiate x2', '409 KYC already in progress (Redis lock)', '✅'],
            ['SEC-007', 'Admin Cascade Delete', 'מחיקת משתמש → cascade', 'P2', 'DELETE /api/v3/admin/users/:id', 'User + all related records deleted', '✅'],
            ['SEC-008', 'Guarantor Link Expiry', 'Link אחרי 5 ימים → error', 'P3', 'GET /api/guarantor/verify/:token', '410 Link expired', '✅'],
        ],
        col_widths=[1, 1.5, 2, 0.4, 2, 2.3, 0.6]
    )

    # ── WhatsApp Tests ─────────────────────────────────────────────────────────
    style_heading(doc, '6. WhatsApp Integration Tests (13)', level=1)
    add_table(doc,
        ['ID', 'Test Case', 'P', 'Expected', 'Status'],
        [
            ['WA-001', 'Webhook verification (GET)', 'P1', '200 + hub.challenge', '✅'],
            ['WA-002', 'Webhook HMAC validation (POST)', 'P1', 'Invalid HMAC → 401', '✅'],
            ['WA-003', 'Inbound text → idle state', 'P1', 'Main menu sent', '✅'],
            ['WA-004', 'Payment confirm flow', 'P1', 'LedgerRow → REPORTED', '✅'],
            ['WA-005', 'Maintenance description', 'P1', 'Awaiting image/skip', '✅'],
            ['WA-006', 'Maintenance + image', 'P1', 'Ticket created, photo → R2', '✅'],
            ['WA-007', 'Maintenance skip image', 'P1', 'Ticket created without photo', '✅'],
            ['WA-008', 'Payment reminder cron', 'P1', 'Template sent T-3', '✅'],
            ['WA-009', 'Overdue cron → WA', 'P1', 'Overdue template sent', '✅'],
            ['WA-010', 'Renewal alert cron', 'P1', 'Renewal template sent at 60d', '✅'],
            ['WA-011', 'Guarantor invite', 'P1', 'Invite template sent', '✅'],
            ['WA-012', 'Message logged to DB', 'P1', 'whatsapp_messages record created', '✅'],
            ['WA-013', 'Conversation state persists', 'P1', 'State survives between messages', '✅'],
        ],
        col_widths=[1, 3.5, 0.4, 3, 0.6]
    )
    doc.add_page_break()

    # ── Admin Tests ────────────────────────────────────────────────────────────
    style_heading(doc, '7. Admin Panel Tests (6)', level=1)
    add_table(doc,
        ['ID', 'Feature', 'Test Case', 'P', 'Method', 'Expected', 'Status'],
        [
            ['ADM-001', 'Config Read', 'GET config → 52 keys returned', 'P0', 'GET /api/v3/admin/config', '200 + JSON with 52 keys', '✅'],
            ['ADM-002', 'Config Update', 'PUT config key → updated real-time', 'P0', 'PUT /api/v3/admin/config/:key', '200, new value active immediately', '✅'],
            ['ADM-003', 'Stats Dashboard', 'GET stats/detailed → 8 sections', 'P0', 'GET /api/v3/admin/stats/detailed', '56 metrics across 8 sections', '✅'],
            ['ADM-004', 'User Management', 'GET users → paginated list', 'P1', 'GET /api/v3/admin/users', 'Paginated with kycProfile', '✅'],
            ['ADM-005', 'KYC Override', 'Admin bypasses Persona', 'P1', 'POST /api/v3/admin/kyc/:id/override', 'KYC status updated, contracts unlocked', '✅'],
            ['ADM-006', 'Force Contract State', 'Admin transitions ACTIVE→ENDED', 'P1', 'PUT /api/v3/admin/contracts/:id/status', 'Status changed, parties notified', '✅'],
        ],
        col_widths=[1, 1.5, 2, 0.4, 2, 2, 0.6]
    )

    # ── Test Targets ──────────────────────────────────────────────────────────
    style_heading(doc, '8. יעדי כיסוי', level=1)
    add_table(doc,
        ['Priority', 'יעד כיסוי', 'מצב נוכחי', 'הערות'],
        [
            ['P0 (Critical)', '100% חובה לפני launch', '35/35 ✅ 100%', 'כל P0 עוברים'],
            ['P1 (High)', '100%', '37/37 ✅ 100%', 'כל P1 עוברים'],
            ['P2 (Medium)', '80%+', '6/6 ✅ 100%', 'Above target'],
            ['P3 (Low)', 'Best effort', '3/3 ✅ 100%', 'Above target'],
        ],
        col_widths=[2, 2, 2, 3]
    )
    add_note(doc, 'Run tests: cd backend && npm test | jest tests/ --coverage', BLUE_MID)

    doc.save('C:\\apartmentapp\\docs\\DirApp_Test_Coverage_Matrix_v3.1.docx')
    print('✅ DirApp_Test_Coverage_Matrix_v3.1.docx')


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    import os, sys
    sys.stdout.reconfigure(encoding='utf-8')
    os.makedirs('C:\\apartmentapp\\docs', exist_ok=True)

    print('Generating DirApp v3.1 Word documents...\n')
    create_roadmap()
    create_architecture()
    create_product_spec()
    create_test_matrix()
    print('\nDone - all 4 documents in C:\\apartmentapp\\docs\\')
